from django.conf import settings
from bible_study.models import Scripture, Question, MyAnswer, MyNote, Note
from datetime import date

from io import StringIO, BytesIO
from docx import Document

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.contrib.auth import get_user_model
User = get_user_model()

# My Answers - ANSWERS - Q&A
# My Reflection - REFLECTION - Reflection only
# Questions Only - QUESTIONS  - Questions only
# My Answers and My Reflection - ALL - Q&A + Reflection
# Notes Only - NOTES - Notes only

def BibleStudyExportWord(scripture, user, what):
    output = BytesIO()
    doc = Document()

    # scripture = Scripture.objects.get(pk=1)
    # user = User.objects.get(pk=1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    doc.add_heading(scripture.title, level=0)

    p = scripture.book + " " + scripture.passage
    doc.add_heading(p, level=1)
    doc.add_heading("kv: " + scripture.get_key_verse_text(), level=2)
    doc.add_paragraph()

    if what == "ANSWERS" or what == "ALL":
        p = doc.add_paragraph()
        p.add_run("Questions").bold = True

        questions = Question.objects.filter(scripture = scripture)
        if questions.exists() :
            for question in questions :
                p = doc.add_paragraph() 
                p.add_run(question.question).bold = True
                my_answer = MyAnswer.objects.filter(question = question, user = user).first()
                if my_answer != None :
                    p = doc.add_paragraph(my_answer.answer)
                    # p.add_run(my_answer.answer).bold = False
                else :
                    doc.add_paragraph("No answer")
        else :
            doc.add_paragraph("No questions")

    if what == "REFLECTION" or what == "ALL":

        p = doc.add_paragraph()
        p.add_run("Reflection").bold = True

        reflection = MyNote.objects.filter(scripture = scripture, user = user).first()
        if reflection != None:
            p = doc.add_paragraph(reflection.note) 
        else :
            doc.add_paragraph("No reflection")
        
    if what == "QUESTIONS" :  
        p = doc.add_paragraph()
        p.add_run("Questions").bold = True

        questions = Question.objects.filter(scripture = scripture)
        if questions.exists() :
            for question in questions :
                p = doc.add_paragraph() 
                p.add_run(question.question).bold = True
        else :
            doc.add_paragraph("No questions")

    if what == "NOTES" :  
        p = doc.add_paragraph()
        p.add_run("Notes").bold = True

        notes = Note.objects.filter(scripture = scripture).order_by('seq_nr')
        if notes.exists() :
            for note in notes :
                if note.verse_heading:
                    p = doc.add_paragraph() 
                    p.add_run(note.verse_heading).bold = True

                if note.verse:
                    p = doc.add_paragraph() 
                    p.underline = True
                    p.add_run(note.verse.verse_text).italic = True
                    
                if note.note:
                    p = doc.add_paragraph(note.note) 
                    
        else :
            doc.add_paragraph("No Notes")

    doc.save(output)
    docx_data = output.getvalue()

    return docx_data



